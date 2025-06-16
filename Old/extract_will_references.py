import re
import pandas as pd
from docx import Document

# Load the document
doc = Document("Copy of dafi36-2110.docx")

# Regex to find the word 'will'
will_pattern = re.compile(r"\bwill\b", re.IGNORECASE)

# Regex to detect section numbers like 1.2, 2.1.3, etc.
section_number_pattern = re.compile(r"^\d+(\.\d+)*")

results = []

# Iterate through paragraphs
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue

    matches = will_pattern.findall(text)
    if matches:
        section_number = ""
        section_title = ""

        # Look back for nearest heading with a number
        for j in range(i, -1, -1):
            heading_candidate = doc.paragraphs[j].text.strip()
            if section_number_pattern.match(heading_candidate):
                section_number = section_number_pattern.match(heading_candidate).group()
                section_title = heading_candidate
                break

        results.append({


            "Paragraph Number": i + 1,
            "Matches": len(matches),
            "Section Number": section_number,
            "Section Title": section_title,
            "Text": text
        })

# Save to CSV
df = pd.DataFrame(results)
df.to_csv("3Will_references_with_section_numbers.csv", index=False)

print("✅ Done! Saved to 'will_references_with_section_numbers.csv'")
